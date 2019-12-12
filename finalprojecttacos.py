import docx  # This will allow me to make the word doc

import requests  # Requests will eventually be used to get the recipes from a url

from docx.shared import Inches  # Used this link: https://python-docx.readthedocs.io/en/latest/ which eventuall will
# help me put the resized taco image in a 6x4 format on line 17

from docx.enum.text import \
    WD_BREAK  # Used this link as a reference on

# https://python-docx.readthedocs.io/en/latest/api/text.html#run-objects

document = docx.Document()

document.add_heading('Random Taco Cookbook', 2)  # This is the header text and the number besides is 2 so it's
# Heading 2

document.add_picture('resized-taco.jpg', width=Inches(6),
                     height=Inches(4))  # This will add the resized taco picture to the word doc in a 6x4 format

document.add_paragraph('The original taco image is owned by Tai\'s Capture')  # This will show the author's name of the
# original taco image

document.add_paragraph(' https://taco-1150.herokuapp.com/random/?full_taco=true')  # The API link

document.add_paragraph('Final project by Devon Reed')  # This is my name

url = 'https://taco-1150.herokuapp.com/random/?full_taco=true'  # this will give a random recipe for the Tai Taco
random_taco_titles = ['First', 'Second', 'Third']  # This will define the variable for first, second and third taco
# recipe that is given to us and will be later used as a for loop

# In line 33,
for i in range(3):
    document.add_heading(f'{random_taco_titles[i]} Taco Recipe', 0)  # add the heading on each random taco recipe:

    response = requests.get(url).json()  # The request() call that will give us three random tacos to use

    seasoning = response['seasoning']['recipe']  # get the seasoning's recipe
    seasoning_name = response['seasoning']['name']  # get the seasoning name

    condiment = response['condiment']['recipe']  # get the condiment's recipe
    condiment_name = response['condiment']['name']  # get the condiment'name

    mixin = response['mixin']['recipe']  # get the mixin's recipe
    mixin_name = response['mixin']['name']  # get the mixin's name

    base_layer = response['base_layer']['recipe']  # get the base-layer's recipe
    base_layer_name = response['base_layer']['name']  # get the base-layer's name

    shell = response['shell']['recipe']  # get the shell's recipe
    shell_name = response['shell']['name']  # get the shell's name

# Add the name of each component of the taco  with 'Heading' style 3 in the lines below of this comment

    document.add_paragraph(seasoning_name, 'Heading 2')  # add the name of the seasoning with heading style 3
    document.add_paragraph(seasoning)  # add the seasoning info

    document.add_paragraph(condiment_name, 'Heading 3')  # add the name of the condiment with heading style 3
    document.add_paragraph(condiment)  # add the condiment info

    document.add_paragraph(mixin_name, 'Heading 3')  # add the name of the mixin with heading style 3
    document.add_paragraph(mixin)  # add the mixin info

    document.add_paragraph(base_layer_name, 'Heading 3')  # add the name of the base-layer in heading style 3
    document.add_paragraph(base_layer)  # add the base-layer info

    document.add_paragraph(shell_name, 'Heading 3')  # add the name of the shell as a heading
    paragraph = document.add_paragraph(shell)  # add the shell type and
    # used this paragraph as reference to add a page break after each random taco recipe was retrieved from line 28
    run = paragraph.add_run()  # add run to able to add the page break
    run.add_break(WD_BREAK.PAGE)  # insert the page break before saving the word doc

document.save('RandomTacoRecipeBook.docx')  # save the word doc before it gets turned in
